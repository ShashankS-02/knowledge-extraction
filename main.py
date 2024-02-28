#  OPEN AI API KEY ---  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
# import openai
# import docx2txt
# from openai import OpenAI
# import json
#
# # filelist = ["44.docx", "45.docx", "46.docx", "03 (2).docx"]
# filelist = ["45.docx"]
# client = OpenAI(api_key="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")
#
#
# def get_para_embeddings(text_to_embed):
#     text_to_embed = text_to_embed.strip().replace("\n", " ").replace("\r", "")
#     response = client.embeddings.create(
#         model="text-embedding-ada-002",
#         input=[text_to_embed],
#         encoding_format="float"
#     )
#     return response.data[0].embedding
#
#
# # paragraph = "My name is Shashank Shandilya. I am studying in NIE."
# # res = get_para_embeddings(paragraph)
# # out_file = open("./Output_vector_files/testfile.json", "w")
# # json_string = json.dump(res, out_file)
# # out_file.close()
#
# i = 1
# for file in filelist:
#     path = "/Users/shashankshandilya/PycharmProjects/Information-Extraction/Final copy/" + file
#     text = docx2txt.process(path)
#     paragraphs = text.split("\n")
#     j = 1
#
#     for paragraph in paragraphs:
#         try:
#             embedding = get_para_embeddings(str(paragraph))
#             out_file = open("/Users/shashankshandilya/PycharmProjects/Information-Extraction/Output_vector_files/Doc" + str(i) + "para" + str(j) + ".json", "w")
#             json.dump(embedding, out_file)
#             out_file.close()
#         except:
#             print(paragraph)
#             print("vector not generated")
#         j += 1
#     i += 1

# ***********************************       FIRST TRY        **********************************************************
# import docx
# import os
# import nltk
# from openai import OpenAI
#
# # nltk.download()
#
# client = OpenAI(api_key="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")
#
# def get_para_embeddings(text_to_embed):
#     text_to_embed = text_to_embed.strip().replace("\n", " ").replace("\r", "")
#     response = client.embeddings.create(
#         model="text-embedding-ada-002",
#         input=[text_to_embed],
#         encoding_format="float"
#     )
#     return response.data[0].embedding
#
# document = docx.Document('/Users/shashankshandilya/PycharmProjects/Information-Extraction/Final copy/51 Sale Deed.docx')
# docText = '\n\n'.join(
#     paragraph.text for paragraph in document.paragraphs
# )
#
# line_file = ''
#
# for line in docText:
#     if line.strip():
#         line_file += line
#
# paragraphs = nltk.tokenize.sent_tokenize(line_file)
#
# for paragraph in paragraphs:
#     embedding = get_para_embeddings(paragraph)
#     print(embedding)

# ******************************************************************************************************************************************

# from openai import OpenAI
# import docx
# import os
#
#
# # filelist = ["44.docx", "45.docx", "46.docx", "03 (2).docx"]
# filelist = ["45.docx"]
# client = OpenAI(api_key="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")
#
#
# def get_para_embeddings(text_to_embed):
#     text_to_embed = text_to_embed.strip().replace("\n", " ").replace("\r", "")
#     response = client.embeddings.create(
#         model="text-embedding-ada-002",
#         input=[text_to_embed],
#         encoding_format="float"
#     )
#     return response.data[0].embedding
#
#
# # paragraph = "My name is Shashank Shandilya. I am studying in NIE."
# # res = get_para_embeddings(paragraph)
# # out_file = open("./Output_vector_files/testfile.json", "w")
# # json_string = json.dump(res, out_file)
# # out_file.close()
#
# for filename in filelist:
#     document = docx.Document('/Users/shashankshandilya/PycharmProjects/Information-Extraction/Final copy/' + filename)
#     docText = '\n\n'.join(
#         paragraph.text for paragraph in document.paragraphs
#     )
#     compressed_file = ''
#
#     compressed_file = os.linesep.join(
#         line
#         for line in docText.splitlines()
#         if line
#     )
#
#     embedding = get_para_embeddings(compressed_file)
#     print(embedding)



import openai
from openai import OpenAI
from docx import Document
import json

client = OpenAI(api_key="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")


def get_para_embeddings(text_to_embed):
    text_to_embed = text_to_embed
    response = client.embeddings.create(
        model="text-embedding-ada-002",
        input=[text_to_embed],
        encoding_format="float"
    )
    return response.data[0].embedding


file_list = ['44.docx']

for file in file_list:
    filename = "/Users/shashankshandilya/PycharmProjects/Information-Extraction/Final copy/" + file
    word_doc = Document(filename)

    out_file = open("/Users/shashankshandilya/PycharmProjects/Information-Extraction/Output_vector_files/" + file +
                    ".json", "w")
    json_data = {'filename': filename, 'paragraphs': []}

    for para in word_doc.paragraphs:
        print(para.text)
        try:
            embedding = get_para_embeddings(para.text)
            json_data['paragraphs'].append({
                'para_text': para.text,
                'embedding': embedding
            })

        except openai.BadRequestError:
            json_data['paragraphs'].append({
                'para_text': para.text,
                'embedding': []
            })

    json.dump(json_data, out_file)
    out_file.close()






